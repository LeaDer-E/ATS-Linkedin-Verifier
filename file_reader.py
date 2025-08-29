# file_reader.py
import os
import re
from docx import Document

def get_all_docx_files(folder_path):
    """
    Return a list of all DOCX files in the folder including subfolders.
    """
    docx_files = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".docx"):
                docx_files.append(os.path.join(root, file))
    return docx_files

def extract_linkedin_links(file_path):
    """
    Extract all LinkedIn links from DOCX file (text + hyperlinks).
    Handles links without http/https and split runs.
    """
    links = set()
    try:
        doc = Document(file_path)

        # Process paragraphs
        pattern = re.compile(
            r"(https?://[^\s]*linkedin\.com[^\s]*|www\.linkedin\.com[^\s]*|linkedin\.com[^\s]+)",
            re.IGNORECASE
        )

        for para in doc.paragraphs:
            text = "".join(run.text for run in para.runs).replace("\n", " ").strip()
            for match in pattern.findall(text):
                clean_link = match.strip()
                if not clean_link.startswith("http"):
                    clean_link = "https://" + clean_link
                links.add(clean_link)

        # Process hyperlinks
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype.lower() and "linkedin.com" in rel.target_ref.lower():
                links.add(rel.target_ref.strip())

    except Exception as e:
        print(f"Error reading {file_path}: {e}")

    return list(links)

    

def extract_phone_number(file_path):
    """
    Extract one valid phone number from the first 10 non-empty lines (including runs) of a DOCX file.
    If not found, return "N/A".
    """
    phone_pattern = re.compile(
        r"(\+?\d{1,4}[\s-]?\(?\d+\)?[\s-]?\d+[\s-]?\d+)"
    )
    try:
        doc = Document(file_path)

        first_lines = []
        for para in doc.paragraphs:
            # ناخد النص من الـ runs كمان
            full_text = "".join(run.text for run in para.runs).strip()
            if not full_text:
                continue
            first_lines.append(full_text)
            if len(first_lines) >= 10:
                break

        text_block = "\n".join(first_lines)
        matches = phone_pattern.findall(text_block)

        clean_numbers = [m for m in matches if len(re.sub(r"\D", "", m)) >= 7]

        if clean_numbers:
            return clean_numbers[0]
        else:
            return "N/A"
    except Exception as e:
        print(f"Error extracting phone from {file_path}: {e}")

    return "N/A"

